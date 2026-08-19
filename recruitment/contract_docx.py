from io import BytesIO
from zipfile import ZipFile

from docx import Document

from .models import ContractTemplate


def _text(value):
    if value is None:
        return ""
    return str(value).strip()


def _thai_date(value):
    if not value:
        return ""
    return value.strftime("%d/%m/%Y")


def contract_placeholders(application, record):
    candidate = application.candidate
    address_parts = [
        _text(candidate.address),
        _text(candidate.tambon),
        _text(candidate.amphure),
        _text(candidate.province),
        _text(candidate.zip),
    ]
    mapping = {
        "ชื่อไทย": _text(candidate.first_name_th),
        "นามสกุลไทย": _text(candidate.last_name_th),
        "คำนำหน้า": _text(candidate.title_name_th),
        "ชื่ออังกฤษ": _text(candidate.first_name),
        "นามสกุลอังกฤษ": _text(candidate.last_name),
        "คำนำหน้าอังกฤษ": _text(candidate.title_name),
        "ชื่อเล่น": _text(candidate.nickname),
        "เบอร์โทร": _text(candidate.phone_number),
        "อีเมล": _text(candidate.email),
        "เลขบัตร": _text(candidate.idcard),
        "ที่อยู่": " ".join(part for part in address_parts if part),
        "แขวง": _text(candidate.tambon),
        "เขต": _text(candidate.amphure),
        "จังหวัด": _text(candidate.province),
        "รหัสไปรษณีย์": _text(candidate.zip),
        "ตำแหน่ง": _text(application.position.title if application.position_id else ""),
        "บริษัท": _text(record.company),
        "สถานที่": _text(record.location),
        "ระดับ": _text(record.employee_level),
        "วันเริ่มงาน": _thai_date(record.start_date),
        "รหัสพนักงาน": _text(record.employee_code),
        "รหัสเอกสาร": _text(record.document_code),
        "ประเภทสัญญา": _text(record.contract_type),
        "first_name_th": _text(candidate.first_name_th),
        "last_name_th": _text(candidate.last_name_th),
        "title_name_th": _text(candidate.title_name_th),
        "first_name": _text(candidate.first_name),
        "last_name": _text(candidate.last_name),
        "title_name": _text(candidate.title_name),
        "phone_number": _text(candidate.phone_number),
        "email": _text(candidate.email),
        "start_date": _thai_date(record.start_date),
        "employee_code": _text(record.employee_code),
        "document_code": _text(record.document_code),
    }
    return {f"{{{{{key}}}}}": value for key, value in mapping.items()}


def _replace_paragraph(paragraph, mapping):
    original = paragraph.text
    if not original:
        return
    updated = original
    for token, value in mapping.items():
        if token in updated:
            updated = updated.replace(token, value)
    if updated == original:
        return
    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated)


def _iter_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs


def _safe_filename(name):
    cleaned = "".join(ch if ch.isalnum() or ch in "._- " else "_" for ch in name).strip()
    return cleaned or "contract"


def fill_template(template_path, mapping):
    document = Document(template_path)
    for paragraph in _iter_paragraphs(document):
        _replace_paragraph(paragraph, mapping)
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def active_templates(contract_type):
    return ContractTemplate.objects.filter(
        contract_type=contract_type,
        is_active=True,
    ).exclude(file="")


def build_contract_download(application, record):
    templates = list(active_templates(record.contract_type))
    if not templates:
        return None, None, "ยังไม่มีแม่แบบสัญญาที่เปิดใช้งานสำหรับประเภทนี้"
    mapping = contract_placeholders(application, record)
    try:
        if len(templates) == 1:
            template = templates[0]
            buffer = fill_template(template.file.path, mapping)
            filename = f"{_safe_filename(record.employee_code)}_{_safe_filename(template.name)}.docx"
            return buffer, filename, None
        archive = BytesIO()
        with ZipFile(archive, "w") as zipped:
            for template in templates:
                filled = fill_template(template.file.path, mapping)
                zipped.writestr(
                    f"{_safe_filename(record.employee_code)}_{_safe_filename(template.name)}.docx",
                    filled.getvalue(),
                )
        archive.seek(0)
        return archive, f"{_safe_filename(record.employee_code)}_contracts.zip", None
    except FileNotFoundError:
        return None, None, "หาไฟล์แม่แบบไม่พบ กรุณาอัปโหลดใหม่อีกครั้ง"


CONTRACT_PLACEHOLDER_HELP = (
    "ใส่รหัสในไฟล์ Word เช่น {{ชื่อไทย}} {{นามสกุลไทย}} {{คำนำหน้า}} "
    "{{ตำแหน่ง}} {{บริษัท}} {{สถานที่}} {{วันเริ่มงาน}} {{รหัสพนักงาน}} {{รหัสเอกสาร}} {{ที่อยู่}}"
)
